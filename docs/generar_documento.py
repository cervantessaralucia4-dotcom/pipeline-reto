from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Estilos ───────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    s = doc.styles[f'Heading {level}']
    s.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)

# ── Portada ───────────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('HealthAnalytics IPS')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Plataforma Inteligente de Analítica Clínica')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documentación Técnica y Manual de Usuario')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Reto Técnico — FullStack + Data Analytics + ETL + Machine Learning')
run.font.size = Pt(12)
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Autora: Sara Lucía Cervantes')
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Tecnólogo en Análisis y Desarrollo de Software — SENA')
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Junio 2026')
run.font.size = Pt(11)

doc.add_page_break()

# ── Tabla de Contenidos ──────────────────────────────────────
doc.add_heading('Tabla de Contenidos', level=1)
toc_items = [
    ('1.', 'Código Fuente', 'Repositorio GitHub'),
    ('2.', 'Base de Datos', 'Script SQL y Modelo Relacional'),
    ('3.', 'Evidencias ETL', 'Logs, Capturas y Reportes'),
    ('4.', 'Documentación Técnica', 'Instalación, Arquitectura, APIs, Dependencias'),
    ('5.', 'Manual de Usuario', 'Login, Dashboard, ETL, Reportes'),
    ('6.', 'Diagramas', 'Arquitectura, Flujo ETL, ERD'),
    ('7.', 'Evidencia Machine Learning', 'Dataset, Métricas y Resultados'),
]
for num, title, desc in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num} {title}')
    run.bold = True
    run.font.size = Pt(12)
    p.add_run(f'\n    {desc}').font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  1. CÓDIGO FUENTE
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. Código Fuente', level=1)
doc.add_paragraph('El código fuente completo del proyecto se encuentra alojado en GitHub:')
p = doc.add_paragraph()
run = p.add_run('https://github.com/cervantessaralucia4-dotcom/pipeline-reto')
run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
run.underline = True

doc.add_paragraph('Estructura del repositorio:')
structure = [
    'backend/ — Configuración Django (settings, urls, wsgi)',
    'patients/ — App de pacientes (CRUD, KPIs, exportación)',
    'etl/ — Pipeline ETL completo (extract, transform, load)',
    'ml/ — Módulo Machine Learning (entrenamiento, predicción)',
    'authentication/ — Autenticación JWT y perfiles por rol',
    'analytics/ — Analítica estadística y segmentación',
    'dashboard/ — Endpoints para dashboard',
    'frontend/ — Aplicación React 19 (dashboard, gráficas, formularios)',
    'datasets/ — Dataset original (.xlsx) y dataset limpio (.csv)',
    'docs/ — Documentación técnica y evidencias',
]
for item in structure:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  2. BASE DE DATOS
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. Base de Datos', level=1)

doc.add_heading('2.1 Script SQL', level=2)
doc.add_paragraph('Motor: PostgreSQL (Neon Cloud). El script completo se encuentra en docs/sql/schema.sql')
doc.add_paragraph('A continuación se presentan las sentencias CREATE TABLE principales:')

sql_blocks = [
    ('Usuarios (auth_user)', '''CREATE TABLE auth_user (
    id              SERIAL PRIMARY KEY,
    password        VARCHAR(128) NOT NULL,
    username        VARCHAR(150) NOT NULL UNIQUE,
    first_name      VARCHAR(150) NOT NULL DEFAULT '',
    last_name       VARCHAR(150) NOT NULL DEFAULT '',
    email           VARCHAR(254) NOT NULL DEFAULT '',
    is_superuser    BOOLEAN NOT NULL DEFAULT FALSE,
    is_staff        BOOLEAN NOT NULL DEFAULT FALSE,
    is_active       BOOLEAN NOT NULL DEFAULT TRUE,
    date_joined     TIMESTAMPTZ NOT NULL
);'''),
    ('Perfiles (authentication_userprofile)', '''CREATE TABLE authentication_userprofile (
    id          SERIAL PRIMARY KEY,
    user_id     INTEGER NOT NULL UNIQUE REFERENCES auth_user(id) ON DELETE CASCADE,
    rol         VARCHAR(20) NOT NULL DEFAULT 'medico'
                CHECK (rol IN ('administrador','medico','analista')),
    created_at  TIMESTAMPTZ NOT NULL DEFAULT NOW()
);'''),
    ('Pacientes (patients_patient)', '''CREATE TABLE patients_patient (
    id                      SERIAL PRIMARY KEY,
    first_name              VARCHAR(100) NOT NULL,
    last_name               VARCHAR(100) NOT NULL,
    age                     INTEGER NOT NULL,
    sex                     VARCHAR(1) NOT NULL CHECK (sex IN ('M','F')),
    weight                  DOUBLE PRECISION NOT NULL,
    height                  DOUBLE PRECISION NOT NULL,
    bmi                     DOUBLE PRECISION NULL,
    systolic_pressure       INTEGER NOT NULL,
    diastolic_pressure      INTEGER NOT NULL,
    heart_rate              INTEGER NOT NULL,
    glucose                 DOUBLE PRECISION NOT NULL,
    cholesterol             DOUBLE PRECISION NOT NULL,
    oxygen_saturation       DOUBLE PRECISION NOT NULL,
    temperature             DOUBLE PRECISION NOT NULL,
    family_history          BOOLEAN NOT NULL DEFAULT FALSE,
    smoker                  BOOLEAN NOT NULL DEFAULT FALSE,
    alcohol_consumption     BOOLEAN NOT NULL DEFAULT FALSE,
    physical_activity       VARCHAR(20) NOT NULL,
    preliminary_diagnosis   VARCHAR(255) NOT NULL,
    disease_risk            VARCHAR(20) NOT NULL,
    consultation_date       DATE NOT NULL,
    created_at              TIMESTAMPTZ NOT NULL DEFAULT NOW()
);'''),
    ('Logs ETL (etl_etllog)', '''CREATE TABLE etl_etllog (
    id                      SERIAL PRIMARY KEY,
    usuario_id              INTEGER NULL REFERENCES auth_user(id) ON DELETE SET NULL,
    fecha_inicio            TIMESTAMPTZ NOT NULL DEFAULT NOW(),
    fecha_fin               TIMESTAMPTZ NULL,
    tiempo_ejecucion        DOUBLE PRECISION NULL,
    registros_extraidos     INTEGER NOT NULL DEFAULT 0,
    registros_duplicados    INTEGER NOT NULL DEFAULT 0,
    registros_nulos         INTEGER NOT NULL DEFAULT 0,
    registros_fuera_rango   INTEGER NOT NULL DEFAULT 0,
    registros_cargados      INTEGER NOT NULL DEFAULT 0,
    estado                  VARCHAR(20) NOT NULL DEFAULT 'en_proceso',
    mensaje                 TEXT NOT NULL DEFAULT '',
    archivo_fuente          VARCHAR(255) NOT NULL DEFAULT ''
);'''),
    ('Métricas ML (ml_mlmetrics)', '''CREATE TABLE ml_mlmetrics (
    id                      SERIAL PRIMARY KEY,
    fecha_entrenamiento     TIMESTAMPTZ NOT NULL DEFAULT NOW(),
    usuario_id              INTEGER NULL REFERENCES auth_user(id) ON DELETE SET NULL,
    accuracy                DOUBLE PRECISION NOT NULL,
    precision               DOUBLE PRECISION NOT NULL,
    recall                  DOUBLE PRECISION NOT NULL,
    f1_score                DOUBLE PRECISION NOT NULL,
    confusion_matrix        JSONB NOT NULL,
    total_registros         INTEGER NOT NULL DEFAULT 0,
    registros_train         INTEGER NOT NULL DEFAULT 0,
    registros_test          INTEGER NOT NULL DEFAULT 0,
    modelo                  VARCHAR(100) NOT NULL DEFAULT 'RandomForestClassifier',
    features                JSONB NOT NULL DEFAULT '[]',
    importancia             JSONB NOT NULL DEFAULT '{}'
);'''),
]
for title, sql in sql_blocks:
    doc.add_heading(title, level=3)
    p = doc.add_paragraph()
    run = p.add_run(sql)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

doc.add_heading('2.2 Modelo Relacional', level=2)
doc.add_paragraph('El modelo consta de 5 tablas principales con las siguientes relaciones:')

rels = [
    ('authentication_userprofile → auth_user', 'Relación 1:1 — Cada usuario tiene un perfil con rol'),
    ('etl_etllog → auth_user', 'Relación N:1 — Un usuario ejecuta múltiples procesos ETL'),
    ('ml_mlmetrics → auth_user', 'Relación N:1 — Un usuario entrena múltiples modelos'),
]
for rel, desc in rels:
    p = doc.add_paragraph()
    run = p.add_run(f'{rel}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  3. EVIDENCIAS ETL
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. Evidencias ETL', level=1)

doc.add_heading('3.1 Dataset Original', level=2)
doc.add_paragraph('Archivo: datasets/dataset_clinico_etl_1800_registros.xlsx')
doc.add_paragraph('Dataset clínico con 1,850 registros que contiene inconsistencias reales:')
issues = [
    '48 registros duplicados',
    'Valores nulos en presión arterial, frecuencia cardíaca, saturación de oxígeno',
    'Errores ortográficos en diagnósticos ("hipertencion", "diabete")',
    'Edades en formato texto ("Treinta", "treinta y dos")',
    'Sexo con formatos inconsistentes (m, f, Femenino, Masculino, M, F)',
    'Valores fuera de rango clínico',
    'Registros sin nombre',
]
for issue in issues:
    doc.add_paragraph(issue, style='List Bullet')

doc.add_heading('3.2 Pipeline ETL', level=2)
doc.add_paragraph('El pipeline consta de 3 fases:')

phases = [
    ('EXTRACT', 'Carga del archivo Excel (1,850 registros)'),
    ('TRANSFORM', 'Limpieza y transformación de datos:'),
    ('LOAD', 'Carga en PostgreSQL y exportación CSV'),
]
for phase, desc in phases:
    p = doc.add_paragraph()
    run = p.add_run(f'{phase}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph('Detalle de la fase TRANSFORM:')
transforms = [
    'Eliminación de duplicados exactos → -48 registros',
    'Eliminación de registros sin nombre → -10 registros',
    'Conversión de tipos (edad texto → entero)',
    'Corrección de sexo basada en el nombre del paciente → 183 correcciones',
    'Imputación de nulos (media, mediana, moda según la variable)',
    'Validación de rangos clínicos → -10 registros fuera de rango',
    'Normalización de diagnósticos (ortografía, mayúsculas)',
    'Cálculo de IMC (peso / altura²)',
    'Clasificación de riesgo según reglas clínicas',
]
for t in transforms:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('3.3 Log de Ejecución', level=2)
log_data = [
    ('Registros extraídos', '1,850'),
    ('Duplicados eliminados', '48'),
    ('Valores nulos tratados', '183'),
    ('Registros fuera de rango', '10'),
    ('Géneros corregidos', '183'),
    ('Registros cargados en BD', '1,792'),
    ('Tiempo de ejecución', '0.83 s'),
]
table = doc.add_table(rows=len(log_data)+1, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Métrica'
hdr[1].text = 'Valor'
for i, (metric, value) in enumerate(log_data, 1):
    table.rows[i].cells[0].text = metric
    table.rows[i].cells[1].text = value

doc.add_paragraph()
doc.add_paragraph('Estado: EXITOSO')
doc.add_paragraph('Usuario: admin_ips')

doc.add_heading('3.4 Reporte Post-ETL', level=2)
doc.add_paragraph('Distribución de pacientes después del ETL:')
post_etl = [
    ('Total pacientes', '1,792'),
    ('Críticos', '1,043 (58.2%)'),
    ('Riesgo Alto', '419 (23.4%)'),
    ('Riesgo Medio', '198 (11.0%)'),
    ('Riesgo Bajo', '132 (7.4%)'),
    ('Glucosa promedio', '207.61 mg/dL'),
    ('IMC promedio', '29.06'),
    ('Edad promedio', '54.8 años'),
]
table2 = doc.add_table(rows=len(post_etl)+1, cols=2)
table2.style = 'Light Grid Accent 1'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Indicador'
hdr2[1].text = 'Valor'
for i, (k, v) in enumerate(post_etl, 1):
    table2.rows[i].cells[0].text = k
    table2.rows[i].cells[1].text = v

doc.add_heading('3.5 Capturas de Pantalla', level=2)
doc.add_paragraph('Las capturas de pantalla del sistema se encuentran en la carpeta docs/evidencias/')
doc.add_paragraph('Ver docs/evidencias/screenshots.md para la guía completa de 25 capturas recomendadas.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  4. DOCUMENTACIÓN TÉCNICA
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. Documentación Técnica', level=1)

doc.add_heading('4.1 Instalación', level=2)
doc.add_paragraph('Requisitos previos:')
prereqs = [
    'Python 3.12 o superior',
    'Node.js 18+ (para frontend)',
    'PostgreSQL (o cuenta Neon Cloud)',
    'Git',
]
for r in prereqs:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph('Pasos de instalación:')
steps = [
    ('1. Clonar repositorio', 'git clone https://github.com/cervantessaralucia4-dotcom/pipeline-reto.git\ncd pipeline-reto'),
    ('2. Entorno virtual', 'python -m venv venv\n# Windows: venv\\Scripts\\activate\n# Mac/Linux: source venv/bin/activate'),
    ('3. Dependencias backend', 'pip install -r requirements.txt'),
    ('4. Variables de entorno', 'Crear archivo .env con:\nSECRET_KEY=tu_clave\nDEBUG=True\nDB_NAME=neondb\nDB_USER=neondb_owner\nDB_PASSWORD=tu_password\nDB_HOST=tu_host.neon.tech\nDB_PORT=5432'),
    ('5. Migraciones', 'python manage.py migrate'),
    ('6. Frontend', 'cd frontend\nnpm install'),
    ('7. Ejecutar', '# Backend:\npython manage.py runserver\n\n# Frontend (otra terminal):\ncd frontend\nnpm start'),
]
for title, code in steps:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    p2 = doc.add_paragraph()
    run2 = p2.add_run(code)
    run2.font.name = 'Consolas'
    run2.font.size = Pt(9)

doc.add_heading('4.2 Arquitectura', level=2)
doc.add_paragraph('El sistema sigue una arquitectura de 3 capas:')
arch_items = [
    ('Frontend (React 19)', 'Dashboard con 7 secciones, gráficas interactivas (Recharts), iconografía (react-icons), consumo de API vía Axios, polling cada 30s, autenticación JWT con localStorage'),
    ('Backend (Django REST Framework)', '5 módulos (patients, etl, analytics, ml, authentication), APIs REST protegidas con JWT, permisos por rol (administrador, medico, analista), documentación Swagger/OpenAPI automática'),
    ('Base de Datos (PostgreSQL)', '5 tablas principales, índices para búsquedas frecuentes, datos JSONB para métricas ML, Neon Cloud como proveedor cloud'),
]
for layer, desc in arch_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{layer}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('4.3 APIs REST', level=2)
doc.add_paragraph('Todos los endpoints protegidos requieren el header:')
p = doc.add_paragraph()
run = p.add_run('Authorization: Bearer <access_token>')
run.font.name = 'Consolas'

api_sections = [
    ('Autenticación', [
        ('POST /api/token/', 'Login — obtener token JWT'),
        ('POST /api/token/refresh/', 'Renovar access token'),
        ('GET /api/auth/me/', 'Perfil del usuario autenticado'),
        ('POST /api/auth/register/', 'Registrar usuario (Admin)'),
        ('GET /api/auth/users/', 'Listar usuarios (Admin)'),
        ('PUT /api/auth/users/{id}/rol/', 'Cambiar rol (Admin)'),
    ]),
    ('Pacientes', [
        ('GET /api/patients/', 'Listar pacientes'),
        ('POST /api/patients/', 'Crear paciente (Admin)'),
        ('GET /api/patients/{id}/', 'Detalle de paciente'),
        ('PUT /api/patients/{id}/', 'Actualizar paciente (Admin)'),
        ('DELETE /api/patients/{id}/', 'Eliminar paciente (Admin)'),
        ('GET /api/patients/export/csv/', 'Exportar CSV'),
    ]),
    ('Dashboard', [
        ('GET /api/dashboard/kpis/', '7 KPIs principales'),
        ('GET /api/dashboard/charts/', 'Datos para gráficas'),
    ]),
    ('ETL', [
        ('POST /api/etl/run/', 'Ejecutar pipeline ETL'),
        ('GET /api/etl/historial/', 'Historial de ejecuciones'),
        ('GET /api/etl/historial/{id}/', 'Detalle de ejecución'),
    ]),
    ('Analytics', [
        ('GET /api/analytics/kpis/', 'KPIs médicos'),
        ('GET /api/analytics/estadisticas/', 'Estadística descriptiva'),
        ('GET /api/analytics/segmentacion/', 'Segmentación de pacientes'),
        ('GET /api/analytics/criticos/', 'Alertas clínicas'),
        ('GET /api/analytics/pacientes-por-filtro/?filtro=<f>', 'Pacientes por filtro'),
    ]),
    ('Machine Learning', [
        ('POST /api/ml/train/', 'Entrenar modelo'),
        ('GET /api/ml/metrics/', 'Métricas del modelo'),
        ('GET /api/ml/historial/', 'Historial de entrenamientos'),
        ('POST /api/ml/predict/', 'Predicción individual'),
    ]),
]
for section, endpoints in api_sections:
    doc.add_heading(section, level=3)
    for endpoint, desc in endpoints:
        p = doc.add_paragraph()
        run = p.add_run(f'{endpoint}  —  ')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        p.add_run(desc)

doc.add_heading('4.4 Dependencias', level=2)
doc.add_paragraph('Backend (Python):')
backend_deps = [
    ('Django 5.2', 'Framework web'),
    ('Django REST Framework', 'APIs REST'),
    ('SimpleJWT', 'Autenticación JWT'),
    ('Pandas', 'Procesamiento ETL'),
    ('NumPy', 'Cálculos numéricos'),
    ('Scikit-learn', 'Machine Learning'),
    ('psycopg2-binary', 'Conector PostgreSQL'),
    ('django-cors-headers', 'CORS'),
    ('drf-spectacular', 'Swagger/OpenAPI'),
    ('openpyxl', 'Lectura de Excel'),
    ('gunicorn', 'Servidor WSGI producción'),
    ('whitenoise', 'Archivos estáticos'),
    ('python-dotenv', 'Variables de entorno'),
    ('dj-database-url', 'Parseo DATABASE_URL'),
]
table3 = doc.add_table(rows=len(backend_deps)+1, cols=2)
table3.style = 'Light Grid Accent 1'
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Paquete'
hdr3[1].text = 'Uso'
for i, (pkg, use) in enumerate(backend_deps, 1):
    table3.rows[i].cells[0].text = pkg
    table3.rows[i].cells[1].text = use

doc.add_paragraph()
doc.add_paragraph('Frontend (React):')
frontend_deps = [
    ('React 19', 'Framework UI'),
    ('Axios', 'Cliente HTTP'),
    ('Recharts', 'Gráficas'),
    ('react-icons', 'Iconografía'),
    ('react-scripts', 'Build tool'),
    ('Chart.js', 'Gráficas adicionales'),
]
table4 = doc.add_table(rows=len(frontend_deps)+1, cols=2)
table4.style = 'Light Grid Accent 1'
hdr4 = table4.rows[0].cells
hdr4[0].text = 'Paquete'
hdr4[1].text = 'Uso'
for i, (pkg, use) in enumerate(frontend_deps, 1):
    table4.rows[i].cells[0].text = pkg
    table4.rows[i].cells[1].text = use

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  5. MANUAL DE USUARIO
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. Manual de Usuario', level=1)

doc.add_heading('5.1 Roles y Permisos', level=2)
roles_data = [
    ('Sección', 'Administrador', 'Médico', 'Analista'),
    ('Dashboard', '✅', '✅', '✅'),
    ('Pacientes', '✅ CRUD', '✅ Lectura', '✅ Lectura'),
    ('ETL', '✅', '❌', '✅'),
    ('Analytics', '✅', '✅', '✅'),
    ('Machine Learning', '✅', '❌', '✅'),
    ('Reportes', '✅ CSV+PDF', '✅ PDF', '✅ CSV+PDF'),
    ('Usuarios', '✅', '❌', '❌'),
]
table5 = doc.add_table(rows=len(roles_data), cols=4)
table5.style = 'Light Grid Accent 1'
for i, row_data in enumerate(roles_data):
    for j, cell_text in enumerate(row_data):
        table5.rows[i].cells[j].text = cell_text

doc.add_heading('5.2 Login', level=2)
doc.add_paragraph('Usuarios preconfigurados (creados automáticamente al migrar):')
users_table = [
    ('Usuario', 'Contraseña', 'Rol'),
    ('admin_ips', 'Admin2026*', 'Administrador'),
    ('medico_ips', 'Medico2026*', 'Médico'),
    ('analista_ips', 'Analista2026*', 'Analista'),
]
table6 = doc.add_table(rows=len(users_table), cols=3)
table6.style = 'Light Grid Accent 1'
for i, row_data in enumerate(users_table):
    for j, cell_text in enumerate(row_data):
        table6.rows[i].cells[j].text = cell_text

doc.add_paragraph()
doc.add_paragraph('Procedimiento:')
login_steps = [
    'Abrir la URL del sistema (http://localhost:3000 en desarrollo, o la URL de Render en producción)',
    'Ingresar usuario y contraseña',
    'Hacer clic en "Ingresar"',
    'El sistema redirige automáticamente al Dashboard principal',
]
for s in login_steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('5.3 Uso del Dashboard', level=2)
doc.add_paragraph('El dashboard se actualiza automáticamente cada 30 segundos y muestra:')
dashboard_items = [
    '7 tarjetas KPI: Total pacientes, Críticos, Riesgo Alto/Medio/Bajo, Glucosa promedio, IMC promedio',
    'Gráfica de pastel (donut): Distribución de riesgo con porcentajes',
    'Gráfica de barras: Pacientes por nivel de riesgo',
    'Indicadores clínicos: Glucosa, IMC, % Críticos, % Riesgo alto',
    'Barra lateral: Navegación entre las 7 secciones del sistema',
]
for item in dashboard_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.4 Uso del ETL', level=2)
etl_steps = [
    'Navegar a la sección "ETL" en la barra lateral',
    'Hacer clic en "Ejecutar ETL ahora"',
    'Esperar a que el pipeline procese los datos (generalmente < 1 segundo)',
    'Revisar las estadísticas: extraídos, duplicados, nulos, fuera de rango, corregidos, cargados',
    'Consultar el historial de ejecuciones en la tabla inferior',
]
for s in etl_steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('5.5 Reportes', level=2)
report_items = [
    'Exportar CSV: Descarga todos los pacientes en formato CSV compatible con Excel',
    'Exportar PDF: Vista optimizada para impresión (oculta sidebar y botones)',
    'El resumen ejecutivo muestra KPIs principales del sistema',
]
for item in report_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  6. DIAGRAMAS
# ══════════════════════════════════════════════════════════════
doc.add_heading('6. Diagramas', level=1)

doc.add_heading('6.1 Diagrama de Arquitectura', level=2)
arch_diagram = '''
┌─────────────────────────────────────────┐
│           Frontend React 19             │
│    Dashboard · KPIs · Gráficas · ML     │
└──────────────────┬──────────────────────┘
                   │  REST API (JWT)
┌──────────────────▼──────────────────────┐
│           Django Backend                │
├─────────────────────────────────────────┤
│  patients/   → CRUD + KPIs             │
│  etl/        → Pipeline ETL + Logs      │
│  analytics/  → Estadísticas + KPIs      │
│  ml/         → Entrenamiento + Predict  │
└──────────────────┬──────────────────────┘
                   │
┌──────────────────▼──────────────────────┐
│       PostgreSQL — Neon Cloud           │
│  patients · etl_logs · ml_metrics      │
└─────────────────────────────────────────┘
'''
p = doc.add_paragraph()
run = p.add_run(arch_diagram)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('6.2 Diagrama de Flujo ETL', level=2)
etl_diagram = '''
dataset_clinico_etl_1800_registros.xlsx
              │
        1. EXTRACT
         └── Leer Excel (1850 registros)
              │
         2. TRANSFORM
          ├── Eliminar duplicados        → -48
          ├── Convertir tipos            (edad texto)
          ├── Tratar nulos               (media/moda)
          ├── Validar rangos clínicos    → -10
          ├── Normalizar sexo            (m/f → M/F)
          ├── Corregir género por nombre → 183
          ├── Normalizar diagnósticos
          ├── Recalcular IMC             (peso/altura²)
          └── Clasificar riesgo          (reglas clínicas)
              │
         3. LOAD
          ├── Guardar en PostgreSQL      (1792 pacientes)
          ├── Exportar CSV limpio
          └── Registrar ETLLog
'''
p = doc.add_paragraph()
run = p.add_run(etl_diagram)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('6.3 Diagrama Entidad-Relación', level=2)
doc.add_paragraph('El modelo relacional consta de 5 tablas:')
erd_tables = [
    'auth_user — Usuarios del sistema (Django)',
    'authentication_userprofile — Perfiles con rol (1:1 con auth_user)',
    'patients_patient — Pacientes clínicos (22 campos)',
    'etl_etllog — Historial de ejecuciones ETL (N:1 con auth_user)',
    'ml_mlmetrics — Métricas de entrenamiento ML (N:1 con auth_user)',
]
for t in erd_tables:
    doc.add_paragraph(t, style='List Bullet')
doc.add_paragraph('Ver docs/sql/erd.md para el diagrama completo con todos los campos y tipos de datos.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  7. EVIDENCIA MACHINE LEARNING
# ══════════════════════════════════════════════════════════════
doc.add_heading('7. Evidencia Machine Learning', level=1)

doc.add_heading('7.1 Dataset de Entrenamiento', level=2)
doc.add_paragraph('Archivo: datasets/clean_clinical_dataset.csv (1792 registros, 22 columnas)')
doc.add_paragraph('Variables predictoras:')
features_data = [
    ('Variable', 'Tipo', 'Descripción', 'Rango'),
    ('edad', 'Numérica', 'Edad del paciente', '18-85'),
    ('IMC', 'Numérica', 'Índice de Masa Corporal', '14.5-42.3'),
    ('glucosa', 'Numérica', 'Glucosa (mg/dL)', '50-400'),
    ('colesterol', 'Numérica', 'Colesterol total (mg/dL)', '100-350'),
    ('presión_sistólica', 'Numérica', 'Presión sistólica (mmHg)', '80-220'),
    ('frecuencia_cardiaca', 'Numérica', 'Frecuencia cardíaca (lpm)', '40-130'),
]
table7 = doc.add_table(rows=len(features_data), cols=4)
table7.style = 'Light Grid Accent 1'
for i, row_data in enumerate(features_data):
    for j, cell_text in enumerate(row_data):
        table7.rows[i].cells[j].text = cell_text

doc.add_paragraph()
doc.add_paragraph('Distribución de clases (variable target: disease_risk):')
class_dist = [
    ('Clase', 'Cantidad', 'Porcentaje'),
    ('Crítico', '1,043', '58.2%'),
    ('Alto', '419', '23.4%'),
    ('Medio', '198', '11.0%'),
    ('Bajo', '132', '7.4%'),
]
table8 = doc.add_table(rows=len(class_dist), cols=3)
table8.style = 'Light Grid Accent 1'
for i, row_data in enumerate(class_dist):
    for j, cell_text in enumerate(row_data):
        table8.rows[i].cells[j].text = cell_text

doc.add_heading('7.2 Modelo', level=2)
doc.add_paragraph('Algoritmo: Random Forest Classifier (scikit-learn)')
doc.add_paragraph('Configuración: n_estimators=100, random_state=42')
doc.add_paragraph('Split: 80% entrenamiento (1,433) / 20% prueba (359)')

doc.add_heading('7.3 Métricas de Desempeño', level=2)
metrics_data = [
    ('Métrica', 'Valor'),
    ('Accuracy', '66.57%'),
    ('Precision', '66.83%'),
    ('Recall', '66.57%'),
    ('F1 Score', '66.54%'),
]
table9 = doc.add_table(rows=len(metrics_data), cols=2)
table9.style = 'Light Grid Accent 1'
for i, row_data in enumerate(metrics_data):
    for j, cell_text in enumerate(row_data):
        table9.rows[i].cells[j].text = cell_text

doc.add_paragraph()
doc.add_paragraph('Importancia de variables:')
importance_data = [
    ('Variable', 'Importancia'),
    ('glucosa', '33.8%'),
    ('presión_sistólica', '25.8%'),
    ('IMC', '10.9%'),
    ('colesterol', '10.6%'),
    ('frecuencia_cardiaca', '9.9%'),
    ('edad', '8.9%'),
]
table10 = doc.add_table(rows=len(importance_data), cols=2)
table10.style = 'Light Grid Accent 1'
for i, row_data in enumerate(importance_data):
    for j, cell_text in enumerate(row_data):
        table10.rows[i].cells[j].text = cell_text

doc.add_heading('7.4 Matriz de Confusión', level=2)
cm = [
    ('', 'Bajo', 'Medio', 'Alto', 'Crítico', 'Total'),
    ('Bajo', '8', '3', '2', '3', '16'),
    ('Medio', '2', '17', '6', '8', '33'),
    ('Alto', '2', '8', '57', '21', '88'),
    ('Crítico', '4', '9', '35', '174', '222'),
    ('Total', '16', '37', '100', '206', '359'),
]
table11 = doc.add_table(rows=len(cm), cols=6)
table11.style = 'Light Grid Accent 1'
for i, row_data in enumerate(cm):
    for j, cell_text in enumerate(row_data):
        table11.rows[i].cells[j].text = cell_text

doc.add_heading('7.5 Ejemplos de Predicción', level=2)
doc.add_paragraph('Ejemplo 1 — Paciente Crítico:')
p = doc.add_paragraph()
run = p.add_run('Input: edad=72, IMC=35.2, glucosa=380.5, colesterol=265, presión=195, frecuencia=102')
run.font.name = 'Consolas'
run.font.size = Pt(9)
p = doc.add_paragraph()
run = p.add_run('Output: riesgo_predicho="Crítico", probabilidad=95%')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph('Ejemplo 2 — Paciente Bajo Riesgo:')
p = doc.add_paragraph()
run = p.add_run('Input: edad=28, IMC=21.5, glucosa=95, colesterol=160, presión=110, frecuencia=72')
run.font.name = 'Consolas'
run.font.size = Pt(9)
p = doc.add_paragraph()
run = p.add_run('Output: riesgo_predicho="Bajo", probabilidad=82%')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('7.6 Archivos Generados', level=2)
files_ml = [
    ('ml/risk_model.pkl', 'Modelo RandomForest serializado', '~4.7 MB'),
    ('ml/label_encoder.pkl', 'LabelEncoder para clases', '~506 B'),
    ('ml/__init__.py', 'Inicializador del módulo', '—'),
]
table12 = doc.add_table(rows=len(files_ml)+1, cols=3)
table12.style = 'Light Grid Accent 1'
hdr12 = table12.rows[0].cells
hdr12[0].text = 'Archivo'
hdr12[1].text = 'Descripción'
hdr12[2].text = 'Tamaño'
for i, (file, desc, size) in enumerate(files_ml, 1):
    table12.rows[i].cells[0].text = file
    table12.rows[i].cells[1].text = desc
    table12.rows[i].cells[2].text = size

# ── Guardar ───────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), 'HealthAnalytics_IPS_Documentacion.docx')
doc.save(output_path)
print(f'Documento guardado en: {output_path}')
